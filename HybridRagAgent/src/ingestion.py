"""
Document Ingestion
==================
Handles PDF, DOCX, and plain-text files → Document chunks.
"""

from __future__ import annotations

import hashlib
import os
from pathlib import Path
from typing import List

from src.rag_engine import Document, split_text


def _hash(text: str, prefix: str = "") -> str:
    return prefix + hashlib.md5(text.encode()).hexdigest()[:8]


def load_pdf(path: str) -> List[Document]:
    import PyPDF2

    docs = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            for j, chunk in enumerate(split_text(text)):
                docs.append(
                    Document(
                        id=_hash(chunk, f"pdf-p{i}-c{j}-"),
                        text=chunk,
                        metadata={"source": Path(path).name, "page": i + 1, "chunk": j},
                    )
                )
    return docs


def load_docx(path: str) -> List[Document]:
    from docx import Document as DocxDocument

    docx = DocxDocument(path)
    full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
    docs = []
    for j, chunk in enumerate(split_text(full_text)):
        docs.append(
            Document(
                id=_hash(chunk, f"docx-c{j}-"),
                text=chunk,
                metadata={"source": Path(path).name, "chunk": j},
            )
        )
    return docs


def load_txt(path: str) -> List[Document]:
    text = Path(path).read_text(encoding="utf-8", errors="ignore")
    docs = []
    for j, chunk in enumerate(split_text(text)):
        docs.append(
            Document(
                id=_hash(chunk, f"txt-c{j}-"),
                text=chunk,
                metadata={"source": Path(path).name, "chunk": j},
            )
        )
    return docs


def ingest_file(path: str) -> List[Document]:
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(path)
    elif ext in (".docx", ".doc"):
        return load_docx(path)
    elif ext in (".txt", ".md"):
        return load_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def ingest_folder(folder: str) -> List[Document]:
    docs = []
    for fp in Path(folder).iterdir():
        if fp.suffix.lower() in (".pdf", ".docx", ".doc", ".txt", ".md"):
            try:
                docs.extend(ingest_file(str(fp)))
                print(f"  ✓ Ingested: {fp.name}")
            except Exception as e:
                print(f"  ✗ Failed {fp.name}: {e}")
    return docs
